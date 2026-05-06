import io
import re
from pathlib import Path
from typing import Any, Dict, List

from bs4 import BeautifulSoup
from docx import Document
from markdownify import markdownify as md
from pypdf import PdfReader


SECTION_RE = re.compile(r"^\s{0,3}(#+\s+.+|[A-ZÇĞİÖŞÜ0-9 _-]{4,}|[0-9]+\.\s+.+)$", re.MULTILINE)
REQ_HINTS = [
    "shall",
    "must",
    "should",
    "required",
    "acceptance criteria",
    "gerekmektedir",
    "gerekir",
    "olmalı",
    "zorunlu",
    "kullanıcı",
    "sistem",
]
FLOW_HINTS = ["→", "->", "navigates", "redirect", "yönlenir", "açılır", "gider"]


def extract_text_from_upload(filename: str, content_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix in {".txt", ".md"}:
        return content_bytes.decode("utf-8", errors="ignore")

    if suffix in {".html", ".htm"}:
        soup = BeautifulSoup(content_bytes.decode("utf-8", errors="ignore"), "html.parser")
        return soup.get_text("\n", strip=True)

    if suffix == ".docx":
        return _extract_docx_text(content_bytes)

    if suffix == ".pdf":
        return _extract_pdf_text(content_bytes)

    return content_bytes.decode("utf-8", errors="ignore")


def build_analysis_doc_context(text: str, filename: str = "") -> Dict[str, Any]:
    normalized = normalize_text(text)

    sections = detect_sections(normalized)
    requirement_lines = detect_requirement_lines(normalized)
    business_rule_lines = detect_business_rule_lines(normalized)
    flow_lines = detect_flow_lines(normalized)
    user_roles = detect_user_roles(normalized)

    return {
        "source": "analysis_document",
        "filename": filename,
        "screen_name": filename or "Analiz Dokümanı",
        "text": normalized,
        "detected_sections": sections,
        "detected_requirements": requirement_lines,
        "detected_business_rules": business_rule_lines,
        "detected_flows": flow_lines,
        "detected_user_roles": user_roles,
        "summary": {
            "char_count": len(normalized),
            "line_count": len([x for x in normalized.splitlines() if x.strip()]),
            "section_count": len(sections),
            "requirement_count": len(requirement_lines),
            "business_rule_count": len(business_rule_lines),
            "flow_count": len(flow_lines),
        },
        "instructions": [
            "Analiz dokümanındaki açık gereksinimleri test case'e dönüştür.",
            "Belirsiz noktaları open_questions altında belirt.",
            "Gereksinim ve iş kuralları çelişirse dokümandaki daha açık ifadeyi tercih et.",
        ],
    }


def normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in (text or "").splitlines()]
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def detect_sections(text: str) -> List[str]:
    sections = []

    for match in SECTION_RE.findall(text):
        clean = match.strip().lstrip("#").strip()
        if clean and clean not in sections:
            sections.append(clean)

    return sections[:80]


def detect_requirement_lines(text: str) -> List[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    found = []

    for line in lines:
        lowered = line.lower()
        if any(hint in lowered for hint in REQ_HINTS):
            found.append(line)
        elif re.match(r"^[-*]\s+.+", line):
            found.append(line)
        elif re.match(r"^[0-9]+\.\s+.+", line):
            found.append(line)

    return dedupe(found)[:200]


def detect_business_rule_lines(text: str) -> List[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    found = []

    rule_keywords = [
        "only",
        "if",
        "when",
        "unless",
        "must",
        "cannot",
        "should not",
        "sadece",
        "ise",
        "olduğunda",
        "yalnızca",
        "zorunlu",
        "izin verilmez",
        "girilmelidir",
    ]

    for line in lines:
        lowered = line.lower()
        if any(keyword in lowered for keyword in rule_keywords):
            found.append(line)

    return dedupe(found)[:150]


def detect_flow_lines(text: str) -> List[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    found = []

    for line in lines:
        lowered = line.lower()
        if any(hint in lowered for hint in FLOW_HINTS):
            found.append(line)

    return dedupe(found)[:120]


def detect_user_roles(text: str) -> List[str]:
    roles = []
    role_keywords = [
        "admin",
        "customer",
        "user",
        "guest",
        "operator",
        "subscriber",
        "müşteri",
        "kullanıcı",
        "misafir",
        "yönetici",
        "admin user",
    ]

    lowered_text = text.lower()

    for role in role_keywords:
        if role in lowered_text:
            roles.append(role)

    return dedupe(roles)


def dedupe(items: List[str]) -> List[str]:
    seen = set()
    result = []

    for item in items:
        key = item.strip().lower()
        if key and key not in seen:
            seen.add(key)
            result.append(item.strip())

    return result


def _extract_docx_text(content_bytes: bytes) -> str:
    document = Document(io.BytesIO(content_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf_text(content_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(content_bytes))
    parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    return "\n\n".join(parts)
